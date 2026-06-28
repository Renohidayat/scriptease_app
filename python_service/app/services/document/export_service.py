import json
from docx import Document
from docx.shared import Pt
from docx2pdf import convert

class ExportService:
    @staticmethod
    def generate_docx(delta_json: str, citations: list, output_path: str):
        doc = Document()
        
        try:
            delta = json.loads(delta_json)
        except Exception:
            delta = {"ops": [{"insert": str(delta_json)}]}
        
        p = doc.add_paragraph()
        ops = delta.get("ops", [])
        
        for op in ops:
            insert = op.get("insert")
            if not isinstance(insert, str):
                continue
                
            attrs = op.get("attributes", {})
            
            # Handle newlines by creating new paragraphs
            lines = insert.split('\n')
            for i, line in enumerate(lines):
                if line:
                    run = p.add_run(line)
                    if attrs.get("bold"):
                        run.bold = True
                    if attrs.get("italic"):
                        run.italic = True
                    if attrs.get("header"):
                        level = attrs["header"]
                        # Approximate header sizes
                        run.font.size = Pt(16 if level == 1 else 14)
                        run.bold = True
                
                if i < len(lines) - 1:
                    p = doc.add_paragraph()
                    
        # Append Daftar Pustaka if citations exist
        if citations:
            doc.add_page_break()
            doc.add_heading("Daftar Pustaka", level=1)
            for cite in citations:
                p = doc.add_paragraph()
                # Use apa_full or just string representation
                apa_text = cite.get('apa_full', str(cite))
                p.add_run(apa_text)
                
        doc.save(output_path)
        return output_path

    @staticmethod
    def generate_pdf_from_docx(docx_path: str, pdf_path: str):
        try:
            convert(docx_path, pdf_path)
            return True
        except Exception as e:
            print(f"Error converting to PDF: {e}")
            return False
